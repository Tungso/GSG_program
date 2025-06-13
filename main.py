import tkinter as tk
from tkinter import messagebox, ttk
from tkcalendar import Calendar
import webbrowser
import os
import sys
import subprocess
from docx import Document
from datetime import datetime, date
import pandas as pd

# ——— 전역 설정 ———
if getattr(sys, 'frozen', False):
    base_dir = sys._MEIPASS
else:
    base_dir = os.path.dirname(os.path.abspath(__file__))

# 학생 목록 로드
student_dir = os.path.join(base_dir, "student_list")
student_file = os.path.join(student_dir, "student_list.xlsx")
try:
    df_students = pd.read_excel(student_file, dtype=str)
    student_names = df_students["학생 이름"].tolist()
except Exception as e:
    # 앱 시작 전에 메시지박스를 호출하면 Tk 인스턴스가 없을 수 있어 임시 root 사용
    temp_root = tk.Tk()
    temp_root.withdraw()
    messagebox.showerror("오류", f"학생 명단을 불러올 수 없습니다:\n{e}")
    temp_root.destroy()
    student_names = []

# 날짜 및 출력 디렉토리 설정
today = datetime.today()
selected_end_year = None
if getattr(sys, 'frozen', False):
    exec_dir = os.path.dirname(sys.executable)
else:
    exec_dir = os.path.dirname(os.path.abspath(__file__))
output_dir = os.path.join(exec_dir, "output")
os.makedirs(output_dir, exist_ok=True)

# ─── 함수 정의 ───

def calculate_days():
    try:
        y = int(entry_start_year.get())
        m = int(entry_start_month.get())
        d = int(entry_start_day.get())
        em = int(entry_end_month.get())
        ed = int(entry_end_day.get())
        ey = selected_end_year or y
        start_date = date(y, m, d)
        end_date = date(ey, em, ed)
        diff = (end_date - start_date).days + 1
        if diff < 1:
            messagebox.showwarning("날짜 오류", "종료 날짜가 시작 날짜보다 이전입니다.")
            entry_days.delete(0, 'end')
            return
        entry_days.delete(0, 'end')
        entry_days.insert(0, str(diff))
    except Exception:
        pass


def show_start_calendar():
    for w in calendar_frame.winfo_children(): w.destroy()
    calendar_frame.grid()
    cal = Calendar(calendar_frame, selectmode="day", year=today.year, month=today.month, day=today.day)
    cal.pack(padx=5, pady=5, expand=True, fill="both")
    def on_select(event):
        sel = cal.selection_get()
        entry_start_year.delete(0, 'end'); entry_start_year.insert(0, sel.year)
        entry_start_month.delete(0, 'end'); entry_start_month.insert(0, sel.month)
        entry_start_day.delete(0, 'end'); entry_start_day.insert(0, sel.day)
        calendar_frame.grid_remove()
        calculate_days()
    cal.bind("<<CalendarSelected>>", on_select)


def show_end_calendar():
    for w in calendar_frame.winfo_children(): w.destroy()
    calendar_frame.grid()
    cal = Calendar(calendar_frame, selectmode="day", year=today.year, month=today.month, day=today.day)
    cal.pack(padx=5, pady=5, expand=True, fill="both")
    def on_select(event):
        sel = cal.selection_get()
        global selected_end_year
        selected_end_year = sel.year
        entry_end_month.delete(0, 'end'); entry_end_month.insert(0, sel.month)
        entry_end_day.delete(0, 'end'); entry_end_day.insert(0, sel.day)
        calendar_frame.grid_remove()
        calculate_days()
    cal.bind("<<CalendarSelected>>", on_select)


def on_student_selected(event):
    name = combo_name.get()
    if name not in student_names: return
    rec = df_students[df_students["학생 이름"] == name].iloc[0]
    entry_grade.delete(0, 'end'); entry_grade.insert(0, rec["학년"])
    entry_class.delete(0, 'end'); entry_class.insert(0, rec["반"])
    entry_number.delete(0, 'end'); entry_number.insert(0, rec["번호"])
    entry_name.delete(0, 'end'); entry_name.insert(0, rec["학생 이름"])
    entry_parent.delete(0, 'end'); entry_parent.insert(0, rec["보호자 이름"])


def update_file_list():
    listbox_files.delete(0, 'end')
    try:
        files = sorted(
            os.listdir(output_dir),
            key=lambda f: os.path.getmtime(os.path.join(output_dir, f)),
            reverse=True
        )
        for f in files:
            listbox_files.insert(tk.END, f)
    except Exception:
        pass


def open_selected_file(event=None):
    sel = listbox_files.curselection()
    if not sel: return
    filename = listbox_files.get(sel[0])
    path = os.path.join(output_dir, filename)
    try:
        if sys.platform.startswith('win'):
            os.startfile(path)
        elif sys.platform.startswith('darwin'):
            subprocess.call(['open', path])
        else:
            subprocess.call(['xdg-open', path])
    except Exception as e:
        messagebox.showerror("열기 오류", f"파일을 열 수 없습니다:\n{e}")


def generate_document():
    # 필드 값 가져오기
    학년   = entry_grade.get()
    반     = entry_class.get()
    번호   = entry_number.get()
    이름   = entry_name.get()
    보호자 = entry_parent.get()
    구분   = var_type.get()
    시작년 = entry_start_year.get()
    시작월 = entry_start_month.get()
    시작일 = entry_start_day.get()
    종료월 = entry_end_month.get()
    종료일 = entry_end_day.get()
    며칠간 = entry_days.get()
    사유   = entry_reason.get("1.0", "end").strip()
    오늘   = datetime.today().strftime("%Y년 %m월 %d일")
    # 파일 경로 설정
    template_path    = os.path.join(base_dir, "extract_template", "template_word.docx")
    filename_docx    = f"결석신고서_{이름}_{시작년.zfill(4)}-{시작월.zfill(2)}-{시작일.zfill(2)}.docx"
    output_path      = os.path.join(output_dir, filename_docx)
    # 템플릿 열기
    try:
        doc = Document(template_path)
    except Exception as e:
        messagebox.showerror("에러", f"템플릿 열기 실패:\n{e}")
        return
    # 치환 맵
    replacements = {
        "{학년}": 학년, "{반}": 반, "{번호}": 번호, "{이름}": 이름,
        "{보호자}": 보호자, "{구분}": 구분,
        "{시작년}": 시작년, "{시작월}": 시작월, "{시작일}": 시작일,
        "{종료월}": 종료월, "{종료일}": 종료일, "{며칠간}": 며칠간,
        "{사유}": 사유, "{오늘날짜}": 오늘,
        "{학생서명}": 이름, "{보호자서명}": 보호자
    }
    # 본문 치환
    for p in doc.paragraphs:
        full = "".join(run.text for run in p.runs)
        new  = full
        for k, v in replacements.items(): new = new.replace(k, v)
        if new != full:
            p.runs[0].text = new
            for run in p.runs[1:]: run.text = ""
    # 테이블 치환
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full = "".join(run.text for run in p.runs)
                    new  = full
                    for k, v in replacements.items(): new = new.replace(k, v)
                    if new != full:
                        p.runs[0].text = new
                        for run in p.runs[1:]: run.text = ""
    # 저장 및 목록 갱신
    try:
        doc.save(output_path)
        messagebox.showinfo("성공", f"{output_path}\n생성 완료!")
        update_file_list()
    except Exception as e:
        with open(os.path.join(exec_dir, "error_log.txt"), "w", encoding="utf-8") as f:
            f.write(str(e))
        messagebox.showerror("저장 실패", f"오류 발생:\n{e}")

# ─── GUI 구성 ───
root = tk.Tk()
root.title("결석계 자동 작성기")
root.geometry("900x700")

# 제목
label_title = tk.Label(root, text="결석계 자동 작성기", font=("맑은 고딕", 16, "bold"))
label_title.grid(row=0, column=0, columnspan=4, pady=(10,20))

# 학생 선택
tk.Label(root, text="학생 선택").grid(row=1, column=0, sticky="e", padx=5, pady=3)
combo_name = ttk.Combobox(root, values=student_names, state="readonly", width=20)
combo_name.grid(row=1, column=1, padx=5, pady=3, sticky="w")
combo_name.set("선택하세요")
combo_name.bind("<<ComboboxSelected>>", on_student_selected)

# 수동 입력 필드
fields = [("학년",2), ("반",3), ("번호",4), ("학생 이름",5), ("보호자 이름",6)]
entries = {}
for label,row in fields:
    tk.Label(root, text=label).grid(row=row, column=0, sticky="e", padx=5, pady=3)
    e = tk.Entry(root)
    e.grid(row=row, column=1, padx=5, pady=3, sticky="w")
    entries[label] = e
entry_grade  = entries["학년"]
entry_class  = entries["반"]
entry_number = entries["번호"]
entry_name   = entries["학생 이름"]
entry_parent = entries["보호자 이름"]

# 결석 구분
tk.Label(root, text="결석 구분").grid(row=7, column=0, sticky="e", padx=5, pady=3)
var_type = tk.StringVar(value="출석인정")
option_menu = tk.OptionMenu(root, var_type, "출석인정", "질병", "기타")
option_menu.config(bg="lightyellow")
option_menu.grid(row=7, column=1, sticky="w", padx=5)

# 시작 날짜
tk.Label(root, text="시작 날짜").grid(row=8, column=0, sticky="e", padx=5, pady=3)
frame_start = tk.Frame(root)
frame_start.grid(row=8, column=1, sticky="w", padx=5)
entry_start_year  = tk.Entry(frame_start, width=5); entry_start_year.pack(side="left")
tk.Label(frame_start, text="년").pack(side="left")
entry_start_month = tk.Entry(frame_start, width=5); entry_start_month.pack(side="left")
tk.Label(frame_start, text="월").pack(side="left")
entry_start_day   = tk.Entry(frame_start, width=5); entry_start_day.pack(side="left")
tk.Label(frame_start, text="일").pack(side="left")
btn_start_cal = tk.Button(root, text="달력", command=show_start_calendar)
btn_start_cal.grid(row=8, column=2, padx=5)

# 종료 날짜
tk.Label(root, text="종료 날짜").grid(row=9, column=0, sticky="e", padx=5, pady=3)
frame_end = tk.Frame(root)
frame_end.grid(row=9, column=1, sticky="w", padx=5)
entry_end_month = tk.Entry(frame_end, width=5); entry_end_month.pack(side="left")
tk.Label(frame_end, text="월").pack(side="left")
entry_end_day   = tk.Entry(frame_end, width=5); entry_end_day.pack(side="left")
tk.Label(frame_end, text="일").pack(side="left")
btn_end_cal = tk.Button(root, text="달력", command=show_end_calendar)
btn_end_cal.grid(row=9, column=2, padx=5)

# 며칠간
tk.Label(root, text="며칠간").grid(row=9, column=3, sticky="e", padx=5)
entry_days = tk.Entry(root, width=5)
entry_days.grid(row=9, column=4, sticky="w", padx=5)

# 사유 입력
tk.Label(root, text="사유").grid(row=10, column=0, sticky="ne", padx=5, pady=3)
entry_reason = tk.Text(root, height=2, width=30)
entry_reason.grid(row=10, column=1, columnspan=2, sticky="w", padx=5, pady=3)

# 생성 버튼
tk.Button(root, text="결석계 생성", command=generate_document, bg="lightgreen").grid(row=11, column=1, pady=10)

# 최근 생성 파일 목록
tk.Label(root, text="최근 생성 파일", font=("맑은 고딕", 12, "bold")).grid(row=12, column=0, padx=5, pady=5, sticky="w")
listbox_files = tk.Listbox(root, width=50, height=8)
listbox_files.grid(row=13, column=0, columnspan=4, padx=10, pady=5, sticky="w")
listbox_files.bind("<Double-Button-1>", open_selected_file)

# 작성 요령 영역
guideline_text = (
    "📌 결석계 작성법\n\n"
    "결석 구분은 인정, 질병, 기타 중 선택합니다.\n\n"
    "사유는 병명, 질환명을 입력합니다.\n"
    "- 생리통\n- 발목염좌\n- 감기\n- 인후염\n\n"
    "달력 버튼을 눌러 날짜를 선택할 수 있습니다.\n"
    "작성 완료 후 '결석계 생성' 버튼을 눌러주세요.\n"
    "출력 후 개인정보 보호를 위해 삭제하세요."
)
guideline_label = tk.Label(root, text=guideline_text, justify="left", anchor="nw",
                            padx=10, pady=10, bg="#f4f4f4", relief="groove",
                            width=40, height=12)
guideline_label.grid(row=1, column=5, rowspan=11, sticky="n", padx=20)

# 달력 표시 프레임
calendar_frame = tk.Frame(root, relief="groove", bd=2)
calendar_frame.grid(row=12, column=5, padx=20, pady=5, sticky="n")
calendar_frame.grid_remove()

# 블로그 링크
def open_blog(event): webbrowser.open("https://blog.naver.com/method917")

copyright_label = tk.Label(root,
    text="© 2025 메쏘드쌤. All rights reserved.",
    font=("맑은 고딕", 9, "underline"), fg="blue", cursor="hand2")
copyright_label.grid(row=14, column=0, columnspan=6, pady=(10,5))
copyright_label.bind("<Button-1>", open_blog)

# 초기 파일 목록 표시
update_file_list()

root.mainloop()
